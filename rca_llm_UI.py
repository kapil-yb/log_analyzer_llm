# rca_generator_with_dropdowns.py
# -----------------------------------------------------------
# This Streamlit app helps generate Root Cause Analysis (RCA)
# reports using LLMs like Mistral, LLaMA3, Dolphin-Mistral via Ollama.
#
# It allows users to upload diagnostic files (PDF, HTML, CSV, etc.),
# extracts content including image OCR, and sends it to the LLM.
# Users can select pre-defined or custom prompts, pick a model,
# and export the response in DOCX or HTML formats.
#
# Best result I have seen is with llamm3 model to generate the RCA. 
# I have tried this code with PDF input for RCA generation, however excel should provide better results. 
#
# ----------------------- HOW TO USE --------------------------
# 1. Install requirements: `pip install -r requirements.txt`
# 2. Make sure Ollama is installed and running locally
#    https://ollama.com/download
# 3. Ensure models are pulled, e.g.:
#    `ollama pull mistral`
#    `ollama pull llama3`
# 4. Run: `streamlit run rca_generator_with_dropdowns.py`
#
# -------------------- requirements.txt ----------------------
# streamlit
# pdfminer.six
# beautifulsoup4
# pandas
# pillow
# pytesseract
# pdf2image
# python-docx
# ollama
# -----------------------------------------------------------

import streamlit as st
import os
import tempfile
from datetime import datetime
import re
from io import BytesIO

from pdfminer.high_level import extract_text as extract_pdf_text
from bs4 import BeautifulSoup
import pandas as pd
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import ollama

st.set_page_config(page_title="RCA Generator", layout="wide")
st.title("🔍 Root Cause Analysis (RCA) Generator")

# --- Upload Section ---
st.markdown("## 📂 Upload Files")
uploaded_files = st.file_uploader(
    "Upload logs, reports, screenshots, or related data (PDF, HTML, CSV, Excel, PNG, JPG)...",
    type=["pdf", "html", "csv", "xlsx", "xls", "png", "jpg", "jpeg"],
    accept_multiple_files=True,
)

# --- Prompt and Model Selection ---
prompt_options = {
    "Generate a RCA report": "You are a senior Site Reliability Engineer assistant. Generate a detailed Root Cause Analysis in the following format:\n\n**Issue Summary**:\n\n**Important Activity timeline from start of the issue till end of the issue i.e. since the issue reported to support**:\n\n**Root Cause**:\n\n**Supporting Evidence from logs / screenshot etc..**:\n\n**Recommended Fix**:\n\n**Preventive Action**:\n",
    "Generate a timeline of the case": "You are a senior product support engineer. Generate a detailed timeline of the case activity, along with steps taken by support, current status and next steps. If there is more than 3 updates in the case for any given day, try to consolidate the update on 6 hourly basis. Typically the support person will be identified by a consistent domain in their email ID and is the assignee of the case.",
    "Summarize the case": "You are a support assistant. Summarize the entire case including issue reported, support actions taken, conclusion and preventive recommendations."
}
prompt_type = st.selectbox("🎯 Select Prompt Type", list(prompt_options.keys()) + ["Custom"])

model_options = {
    "mistral": "Mistral: Balanced speed and quality, great for summaries and RCA.",
    "llama3": "LLaMA 3: Better contextual understanding, best for timeline extraction.",
    "dolphin-mistral": "Dolphin Mistral: Enhanced instructions-following, suitable for most RCA tasks."
}
selected_model = st.selectbox("🧠 Choose LLM Model", list(model_options.keys()), format_func=lambda x: f"{x} - {model_options[x]}")

# --- Additional User Input ---
extra_info = st.text_area("🧾 Additional Manual Context (if available)", "")
custom_prompt = ""
if prompt_type == "Custom":
    custom_prompt = st.text_area("✍️ Enter Your Custom Prompt", height=200, placeholder="e.g., Provide a forensic analysis of the incident with log correlation")

# --- File Content Extraction ---
def extract_content(file):
    extension = os.path.splitext(file.name)[1].lower()

    if extension == ".pdf":
        combined_text = ""
        image_text = ""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file.read())
            pdf_path = tmp.name
            combined_text = extract_pdf_text(pdf_path)
            try:
                images = convert_from_path(pdf_path)
                for image in images:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_tmp:
                        image.save(img_tmp.name, "PNG")
                        image_text += pytesseract.image_to_string(Image.open(img_tmp.name)) + "\n"
            except Exception as e:
                st.warning(f"Image extraction failed: {str(e)}")
        return combined_text + "\n[Extracted OCR Images]\n" + image_text

    elif extension == ".html":
        soup = BeautifulSoup(file.read(), "html.parser")
        return soup.get_text()

    elif extension in [".csv", ".xlsx", ".xls"]:
        return pd.read_csv(file).to_string() if extension == ".csv" else pd.read_excel(file).to_string()

    elif extension in [".png", ".jpg", ".jpeg"]:
        return pytesseract.image_to_string(Image.open(file))

    return ""

# --- RCA Generation ---
conf_check = st.checkbox("✅ I confirm that uploaded content contains no confidential or personal data.")
if st.button("🚀 Generate RCA"):
    if not conf_check:
        st.warning("⚠️ You must confirm that uploaded data contains no confidential or personal information.")
    elif not uploaded_files:
        st.warning("Please upload at least one file.")
    else:
        st.info("⏳ Processing files and generating RCA... Please wait.")
        combined_context = ""
        for f in uploaded_files:
            extracted_text = extract_content(f)
            combined_context += f"\n### File: {f.name}\n{extracted_text}"

        if extra_info.strip():
            combined_context += f"\n### Manual User Input:\n{extra_info.strip()}"

        selected_prompt = custom_prompt if prompt_type == "Custom" else prompt_options[prompt_type]
        full_prompt = f"{selected_prompt}\n\n### Context from Uploaded Files:\n{combined_context}"

        response = ollama.chat(
            model=selected_model,
            messages=[
                {"role": "system", "content": "You are a helpful assistant specialized in Root Cause Analysis."},
                {"role": "user", "content": full_prompt},
            ]
        )

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        full_output = f"{response['message']['content']}\n\n---\n📅 Generated on: {timestamp}"

        st.markdown("### ✅ RCA Report")
        st.markdown(full_output)

        with st.expander("📤 Export RCA Report"):
            export_format = st.selectbox("Choose export format", ["DOCX", "HTML"])
            raw_filename = "rca_report"
            safe_filename = re.sub(r'\W+', '_', raw_filename.lower()).strip('_')
            filename = f"{safe_filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

            if export_format == "DOCX":
                doc = Document()
                doc.add_heading("Root Cause Analysis Report", level=1)
                for line in full_output.split("\n"):
                    if line.strip().startswith("###"):
                        doc.add_heading(line.replace("###", "").strip(), level=2)
                    elif line.strip().startswith("**") and "**" in line[2:]:
                        parts = line.split("**")
                        doc.add_paragraph(f"{parts[1]}: {parts[3] if len(parts) > 3 else ''}")
                    elif line.strip().startswith("-"):
                        doc.add_paragraph(line.strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                st.download_button("⬇️ Download DOCX", data=docx_buffer, file_name=f"{filename}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            elif export_format == "HTML":
                html_content = f"<html><body><h1>Root Cause Analysis</h1><pre>{full_output}</pre></body></html>"
                st.download_button("⬇️ Download HTML", data=html_content, file_name=f"{filename}.html", mime="text/html")
